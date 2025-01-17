import os
from flask import Flask, request, send_file, render_template
from PyPDF2 import PdfMerger
from docx import Document
from io import BytesIO

app = Flask(__name__)

# Route to serve the HTML page
@app.route('/')
def index():
    return render_template('merge.html')

# Route to handle merging documents
@app.route('/merge-documents', methods=['POST'])
def merge_documents():
    files = request.files.getlist('documents[]')
    output_format = request.args.get('format', 'pdf')  # Get output format from query params
    output_filename = request.form.get('filename', 'merged_output')
    
    # Merge PDFs
    if output_format == 'pdf':
        merger = PdfMerger()
        for file in files:
            if file.filename.endswith('.pdf'):
                merger.append(file)
        output_path = f'{output_filename}.pdf'
        with open(output_path, 'wb') as f_out:
            merger.write(f_out)
        return send_file(output_path, as_attachment=True)

    # Merge Word files
    elif output_format == 'word':
        doc = Document()
        for file in files:
            if file.filename.endswith('.docx'):
                temp_doc = Document(file)
                for para in temp_doc.paragraphs:
                    doc.add_paragraph(para.text)
        output_path = f'{output_filename}.docx'
        doc.save(output_path)
        return send_file(output_path, as_attachment=True)

    # Merge text files
    elif output_format == 'text':
        output_text = ''
        for file in files:
            if file.filename.endswith('.txt'):
                output_text += file.read().decode('utf-8') + '\n'
        output_path = f'{output_filename}.txt'
        with open(output_path, 'w') as f_out:
            f_out.write(output_text)
        return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
